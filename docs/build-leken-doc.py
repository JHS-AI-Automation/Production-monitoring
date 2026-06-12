"""Bouwt docx + pdf van een markdown-lekendocument (geen pandoc nodig).

Waarom geen pandoc: pandoc/pypandoc staan op deze machine niet geinstalleerd. Deze
bouwer gebruikt alleen python-docx (docx schrijven) + docx2pdf (docx -> pdf via het
geinstalleerde Word). Hergebruikt het DGS-logo uit internal/assets voor de cover.

Ondersteunt een bewust kleine markdown-subset (genoeg voor dit document):
  # Titel (eerste regel -> cover)        ## sectie        ### subsectie
  gewone alinea's (blanco regel = scheiding)
  - bullets
  | tabel | met | --- | scheidingsrij |
  ```code-blok```  (monospace, regels behouden)
  > callout
  ---  (horizontale lijn -> overgeslagen)
  inline: **vet** en `code`

Gebruik (system python heeft docx + docx2pdf):
  python docs/build-leken-doc.py                          # default: architectuur-voor-leken
  python docs/build-leken-doc.py --src docs/ander-doc.md  # docx+pdf naast de bron
"""
import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve()
REPO_ROOT = HERE.parents[4]          # .../Uland AI
DASH = HERE.parents[1]               # .../alarm-dashboard
LOGO = REPO_ROOT / "internal" / "assets" / "logo-dgs.png"

DEFAULT_SRC = DASH / "docs" / "optimax-architectuur-voor-leken.md"

RED = RGBColor(0xED, 0x1C, 0x24)     # DGS-rood
DARK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x78, 0x78, 0x78)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"
MONO = "Consolas"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
IMG_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)$")


def add_inline(paragraph, text, *, size=11, color=DARK, base_bold=False):
    """Voeg tekst toe met **vet** en `code` als aparte runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.font.name = FONT
            r.font.size = Pt(size)
            r.font.bold = True
            r.font.color.rgb = color
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = MONO
            r.font.size = Pt(size - 1)
            r.font.color.rgb = RGBColor(0x10, 0x40, 0x60)
        else:
            r = paragraph.add_run(part)
            r.font.name = FONT
            r.font.size = Pt(size)
            r.font.bold = base_bold
            r.font.color.rgb = color


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _borders(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "2")
        b.set(qn("w:color"), hex_color)
        borders.append(b)
    tcPr.append(borders)


def add_table(doc, rows):
    """rows = lijst van cel-lijsten; rij 0 = header."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""
            if ri == 0:
                _shade(cell, "ED1C24")
                _borders(cell, "ED1C24")
                add_inline(cell.paragraphs[0], text, size=9, color=WHITE, base_bold=True)
            else:
                _shade(cell, "F8F6F6" if ri % 2 == 0 else "FFFFFF")
                _borders(cell, "C8C8C8")
                add_inline(cell.paragraphs[0], text, size=9, color=DARK,
                           base_bold=(ci == 0))


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = MONO
        run.font.size = Pt(8)
        run.font.color.rgb = DARK
        if i < len(lines) - 1:
            run.add_break()


def add_page_number_footer(doc):
    section = doc.sections[-1]
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GREY
    for kind, text in (("begin", None), ("instr", "PAGE"), ("separate", None),
                       ("text", "1"), ("end", None)):
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        elif kind == "text":
            el = OxmlElement("w:t")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        run._element.append(el)


def build_cover(doc, title, subtitle):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(90)
        p.add_run().add_picture(str(LOGO), width=Mm(55))
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(40)
    r = pt.add_run(title)
    r.font.name = FONT
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RED
    if subtitle:
        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ps.paragraph_format.space_before = Pt(12)
        r = ps.add_run(subtitle)
        r.font.name = FONT
        r.font.size = Pt(13)
        r.font.color.rgb = DARK
    doc.add_page_break()


def heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.bold = True
    r.font.size = Pt(16 if level == 1 else 12)
    r.font.color.rgb = RED if level == 1 else DARK


def add_image(doc, base_dir, alt, rel_path):
    """Afbeelding gecentreerd, paginabreed (160mm past binnen de marges van 22mm)."""
    img = (base_dir / rel_path).resolve()
    if not img.exists():
        print(f"WAARSCHUWING: afbeelding niet gevonden, overgeslagen: {img}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(img), width=Mm(160))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(alt)
        r.font.name = FONT
        r.font.size = Pt(8)
        r.font.italic = True
        r.font.color.rgb = GREY


def parse_and_build(md_text, doc, base_dir):
    lines = md_text.split("\n")
    # Cover: eerste "# " + eerstvolgende niet-lege alinea als subtitle.
    title, subtitle = "Document", ""
    idx = 0
    while idx < len(lines):
        if lines[idx].startswith("# "):
            title = lines[idx][2:].strip()
            idx += 1
            break
        idx += 1
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    if idx < len(lines) and not lines[idx].startswith(("#", "|", "-", ">", "```")):
        subtitle = lines[idx].strip()
        idx += 1
    build_cover(doc, title, subtitle)

    para_buf = []

    def flush_para():
        if not para_buf:
            return
        text = " ".join(s.strip() for s in para_buf).strip()
        para_buf.clear()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, text)

    i = idx
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_para()
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            add_code_block(doc, code)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_para()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):  # sla scheidingsrij over
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        if stripped == "---":
            flush_para()
            i += 1
            continue

        img_match = IMG_RE.match(stripped)
        if img_match:
            flush_para()
            add_image(doc, base_dir, img_match.group(1), img_match.group(2))
            i += 1
            continue

        if stripped.startswith("### "):
            flush_para()
            heading(doc, stripped[4:], level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            flush_para()
            heading(doc, stripped[3:], level=1)
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1  # extra titel-regel: overslaan (cover is al gemaakt)
            continue

        if stripped.startswith("- "):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue

        if stripped.startswith("> "):
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(5)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, stripped[2:], color=GREY)
            for run in p.runs:
                run.font.italic = True
            i += 1
            continue

        if not stripped:
            flush_para()
            i += 1
            continue

        para_buf.append(line)
        i += 1

    flush_para()


def main():
    parser = argparse.ArgumentParser(description="Markdown -> DGS-docx (+pdf via Word)")
    parser.add_argument("--src", type=Path, default=DEFAULT_SRC,
                        help="Bron-markdown; docx en pdf komen ernaast te staan")
    args = parser.parse_args()
    src = args.src.resolve()
    dst_docx = src.with_suffix(".docx")
    dst_pdf = src.with_suffix(".pdf")

    md = src.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    parse_and_build(md, doc, src.parent)
    add_page_number_footer(doc)
    doc.save(str(dst_docx))
    Document(str(dst_docx))  # sanity: heropenen
    print(f"DOCX klaar: {dst_docx} ({dst_docx.stat().st_size // 1024} kB)")

    try:
        from docx2pdf import convert
        convert(str(dst_docx), str(dst_pdf))
        print(f"PDF klaar:  {dst_pdf} ({dst_pdf.stat().st_size // 1024} kB)")
    except Exception as e:
        print(f"PDF overgeslagen ({type(e).__name__}: {e}). DOCX is wel klaar; "
              f"open hem in Word en 'Opslaan als PDF' kan altijd handmatig.")


if __name__ == "__main__":
    main()
